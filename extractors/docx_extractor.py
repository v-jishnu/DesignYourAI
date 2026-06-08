"""
DOCX extractor for MCQ extraction from Word documents.
LLM-first approach with MAANG-level quality validation.
"""

import json
import random
import asyncio
from typing import List
from pathlib import Path
import re

try:
    from docx import Document
except ImportError:
    Document = None

from extractors.base_extractor import BaseExtractor
from config.schemas import MCQ
from utils.text_processor import extract_option_text
from validators.quality_validator import QualityValidator


class DOCXExtractor(BaseExtractor):
    """Extract MCQs from DOCX documents with MAANG-level quality validation."""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.validator = QualityValidator()

    async def extract(self, source: Path) -> List[MCQ]:
        """
        Extract MCQs from DOCX with quality validation and distribution enforcement.

        Flow:
        1. Extract text from paragraphs and tables
        2. Try regex for structured MCQs (fast path)
        3. If not found or quality < 80%, use LLM extraction (smart path)
        4. Validate all MCQs through quality filter
        5. Return validated MCQs

        Args:
            source: Path to DOCX file

        Returns:
            List of quality-validated MCQs
        """
        if Document is None:
            self.log("python-docx not installed. Install with: pip install python-docx", 'error')
            return []

        source = Path(source)
        if not source.exists():
            self.log(f"DOCX file not found: {source}", 'error')
            return []

        self.log(f"Extracting from DOCX: {source.name}")

        try:
            doc = Document(source)

            # Extract text from paragraphs
            full_text = "\n".join([para.text for para in doc.paragraphs])

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    full_text += "\n" + row_text

            if not full_text or len(full_text.strip()) < 50:
                self.log(f"No meaningful text extracted from {source.name}", 'warning')
                return []

            # 1. Quick check: try regex extraction for perfectly structured MCQs (fast, no API call)
            regex_mcqs = self._parse_mcq_text(full_text, str(source))

            if regex_mcqs:
                # Validate regex-extracted MCQs
                passed, failed, stats = self.validator.batch_validate(regex_mcqs)
                self.log(f"Regex extraction: {len(passed)}/{len(regex_mcqs)} passed quality check "
                        f"({stats['pass_rate']*100:.1f}%)")

                # If pass rate is good (>80%), return them
                if stats['pass_rate'] >= 0.8:
                    return passed

                # Otherwise, fall through to LLM extraction for better quality
                self.log("Regex quality insufficient, falling back to LLM extraction for upgrade...")

            # 2. LLM-first approach: send raw text to LLM for intelligent extraction
            if self.llm_client:
                self.log("Using LLM for MAANG-level extraction with quality validation...")
                mcqs = await self._llm_extract(full_text, str(source))
                self.log(f"Extracted {len(mcqs)} quality-validated MCQs from {source.name}")
                return mcqs
            else:
                self.log("No LLM client available for smart extraction", 'warning')
                return []

        except Exception as e:
            self.log(f"Error extracting from DOCX {source}: {e}", 'error')
            return []

    async def _llm_extract(self, text: str, source: str) -> List[MCQ]:
        """
        LLM-based extraction with category distribution and quality validation.

        Important: This handles BOTH:
        1. Extracting existing MCQs/Q&As from source
        2. Upgrading quality to meet MAANG standards
        3. Enforcing 50/25/25 distribution

        Args:
            text: Full extracted text from DOCX
            source: Source file path

        Returns:
            List of quality-validated MCQ objects
        """
        batches = self._split_for_llm(text)
        self.log(f"Split text into {len(batches)} batches for LLM processing")

        all_mcqs = []

        for i, batch in enumerate(batches, 1):
            self.log(f"Processing batch {i}/{len(batches)} ({len(batch)} chars)...")

            # Extract with distribution (50/25/25)
            batch_mcqs = await self._extract_batch_with_distribution(batch, source)

            # Validate quality
            if batch_mcqs:
                passed, failed, stats = self.validator.batch_validate(batch_mcqs)

                self.log(f"Batch {i}: {stats['pass_rate']*100:.1f}% pass rate "
                        f"({len(passed)}/{len(batch_mcqs)} passed)")

                # Log first failure for debugging
                if failed:
                    mcq, reason = failed[0]
                    self.log(f"Quality fail example: {mcq.question_text[:50]}... | Reason: {reason}", 'warning')

                all_mcqs.extend(passed)

        return all_mcqs

    async def _extract_batch_with_distribution(self, text: str, source: str) -> List[MCQ]:
        """
        Extract MCQs from batch with category distribution and quality upgrade.

        This method:
        1. Detects if content has existing MCQs/Q&As or is descriptive
        2. Uses appropriate category-specific prompts
        3. Generates with 50/25/25 distribution
        4. Upgrades quality if needed

        Args:
            text: Text batch (~3500 chars)
            source: Source file path

        Returns:
            List of MCQ objects with distribution enforced
        """
        from classifiers.prompt_templates import get_maang_prompt

        # Estimate how many questions this batch can yield
        estimated_questions = max(3, len(text) // 500)  # ~1 question per 500 chars

        # Calculate distribution (50/25/25)
        num_conceptual = int(estimated_questions * 0.5)
        num_mathematical = int(estimated_questions * 0.25)
        num_application = estimated_questions - num_conceptual - num_mathematical

        all_mcqs = []

        # Generate for each category
        for category, target_count in [
            ('Conceptual', num_conceptual),
            ('Mathematical', num_mathematical),
            ('Application', num_application)
        ]:
            if target_count == 0:
                continue

            # Use category-specific MAANG prompt with few-shot examples
            prompt = get_maang_prompt(category, text)

            # Call LLM with retry
            response = await self._call_llm_with_retry(prompt)
            if not response:
                continue

            # Parse JSON response
            mcqs_data = self._parse_llm_response(response)

            # Limit to target count
            mcqs_data = mcqs_data[:target_count]

            # Create MCQ objects with answer shuffling
            for data in mcqs_data:
                # Ensure category matches what was requested
                data['category'] = category
                mcq = self._create_mcq_with_shuffle(data, source)
                if mcq:
                    all_mcqs.append(mcq)

        return all_mcqs

    async def _call_llm_with_retry(self, prompt: str, max_retries: int = 3) -> str:
        """
        Call LLM with exponential backoff retry for rate limits.
        Supports both Gemini and OpenAI-compatible APIs.

        Args:
            prompt: Prompt to send to LLM
            max_retries: Maximum number of retry attempts

        Returns:
            LLM response text, or None if all retries fail
        """
        for attempt in range(max_retries):
            try:
                # Check if using Gemini or OpenAI-compatible API
                if hasattr(self.llm_client, 'client') and hasattr(self.llm_client.client, 'models'):
                    # Gemini API
                    response = self.llm_client.client.models.generate_content(
                        model=self.llm_client.model_name,
                        contents=prompt
                    )
                    return response.text
                else:
                    # OpenAI-compatible API (Groq, OpenAI, Together)
                    response = await self.llm_client.chat.completions.create(
                        model=self.llm_client.model_name,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.7
                    )
                    return response.choices[0].message.content

            except Exception as e:
                error_str = str(e)
                if '429' in error_str or 'rate_limit' in error_str.lower():
                    wait_time = min(60 * (attempt + 1), 120)
                    self.log(f"Rate limited, waiting {wait_time}s (attempt {attempt + 1}/{max_retries})...")
                    await asyncio.sleep(wait_time)
                else:
                    self.log(f"LLM call error: {e}", 'error')
                    if attempt < max_retries - 1:
                        await asyncio.sleep(5)
                    else:
                        return None

        self.log(f"Failed after {max_retries} retries", 'error')
        return None

    def _create_mcq_with_shuffle(self, data: dict, source: str) -> MCQ:
        """
        Create MCQ from LLM response data with code-level answer randomization.

        Args:
            data: Dictionary with MCQ fields from LLM
            source: Source file path

        Returns:
            MCQ object with shuffled answer positions, or None
        """
        try:
            import uuid
            from datetime import datetime

            options = [data['option_a'], data['option_b'], data['option_c'], data['option_d']]
            correct_letter = data['correct_answer'].strip().upper()

            if correct_letter not in ['A', 'B', 'C', 'D']:
                self.log(f"Invalid correct_answer '{correct_letter}', skipping", 'warning')
                return None

            correct_idx = ord(correct_letter) - ord('A')

            # Shuffle by index so duplicate option text can't mismatch the correct answer
            indices = list(range(4))
            random.shuffle(indices)
            options = [options[i] for i in indices]
            new_correct_idx = indices.index(correct_idx)
            new_correct_letter = chr(ord('A') + new_correct_idx)

            # Validate category
            category = data.get('category', 'Conceptual')
            if category not in ['Mathematical', 'Application', 'Conceptual']:
                category = 'Conceptual'

            mcq = MCQ(
                question_id=str(uuid.uuid4()),
                question_text=data['question_text'].strip(),
                option_a=options[0],
                option_b=options[1],
                option_c=options[2],
                option_d=options[3],
                correct_answer=new_correct_letter,
                explanation=data.get('explanation', ''),
                category=category,
                source=source,
                date_added=datetime.now()
            )

            return mcq

        except Exception as e:
            self.log(f"Error creating MCQ: {e}", 'warning')
            return None

    def _split_for_llm(self, text: str, max_chars: int = 3500) -> List[str]:
        """
        Split text into LLM-friendly batches without breaking content.

        Args:
            text: Full text to split
            max_chars: Maximum characters per batch

        Returns:
            List of text batches
        """
        if len(text) <= max_chars:
            return [text]

        # Split by double newlines (paragraph boundaries)
        paragraphs = text.split('\n\n')

        batches = []
        current_batch = []
        current_length = 0

        for para in paragraphs:
            if current_length + len(para) > max_chars and current_batch:
                batches.append('\n\n'.join(current_batch))
                current_batch = [para]
                current_length = len(para)
            else:
                current_batch.append(para)
                current_length += len(para)

        if current_batch:
            batches.append('\n\n'.join(current_batch))

        return batches

    def _parse_llm_response(self, response_text: str) -> List[dict]:
        """
        Parse LLM JSON response with cleanup.

        Args:
            response_text: Raw LLM response

        Returns:
            List of MCQ dictionaries with validated fields
        """
        try:
            text = response_text.strip()

            # Strip markdown code fences
            if text.startswith('```json'):
                text = text.split('\n', 1)[1] if '\n' in text else text[7:]
            elif text.startswith('```'):
                text = text.split('\n', 1)[1] if '\n' in text else text[3:]
            if text.endswith('```'):
                text = text.rsplit('```', 1)[0]
            text = text.strip()

            # Find JSON array boundaries
            start = text.find('[')
            end = text.rfind(']')
            if start != -1 and end != -1:
                text = text[start:end + 1]

            data = json.loads(text)

            if not isinstance(data, list):
                self.log("LLM response is not a JSON array", 'warning')
                return []

            # Validate each MCQ has required fields
            required = ['question_text', 'option_a', 'option_b', 'option_c', 'option_d', 'correct_answer']
            valid = [item for item in data if all(f in item for f in required)]

            if len(valid) < len(data):
                self.log(f"Filtered {len(data) - len(valid)} invalid MCQs (missing fields)", 'warning')

            return valid

        except json.JSONDecodeError as e:
            self.log(f"Failed to parse LLM JSON response: {e}", 'error')
            return []
        except Exception as e:
            self.log(f"Error parsing LLM response: {e}", 'error')
            return []

    def _parse_mcq_text(self, text: str, source: str) -> List[MCQ]:
        """
        Parse MCQs from extracted text using regex (fast fallback for structured MCQs).
        Only used when text clearly has A) B) C) D) format.
        """
        mcqs = []

        # Check if text has enough MCQ markers to warrant regex extraction
        mcq_markers = ['A)', 'B)', 'C)', 'D)', '(a)', '(b)', '(c)', '(d)']
        marker_count = sum(1 for marker in mcq_markers if marker in text)
        if marker_count < 3:
            return []

        # Pattern: Question with A) B) C) D)
        pattern = r'(?P<question>.+?\?)\s*(?:A\)|a\))\s*(?P<opt_a>.+?)\s*(?:B\)|b\))\s*(?P<opt_b>.+?)\s*(?:C\)|c\))\s*(?P<opt_c>.+?)\s*(?:D\)|d\))\s*(?P<opt_d>.+?)(?:\n|$)'

        for match in re.finditer(pattern, text, re.DOTALL):
            mcq = self._create_mcq(
                question_text=match.group('question').strip(),
                option_a=extract_option_text(match.group('opt_a').strip()),
                option_b=extract_option_text(match.group('opt_b').strip()),
                option_c=extract_option_text(match.group('opt_c').strip()),
                option_d=extract_option_text(match.group('opt_d').strip()),
                source=source
            )

            if mcq:
                mcqs.append(mcq)

        return mcqs
